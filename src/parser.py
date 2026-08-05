"""Extracts raw text from resume files. Supports PDF, DOCX, and TXT.

Scanned/image-only PDFs are a known, deliberate non-goal for this build --
see README "Known limitations". They raise a clear ParseError instead of
silently returning empty text.
"""
from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document


class ParseError(Exception):
    """Raised when a resume file cannot be parsed into usable text."""


def parse_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".txt":
        return _parse_txt(path)
    raise ParseError(f"Unsupported file type '{suffix}' for {path.name}")


def _parse_pdf(path: Path) -> str:
    try:
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        raise ParseError(f"Could not open PDF {path.name}: {e}") from e

    text = "\n".join(text_parts).strip()
    if not text:
        raise ParseError(
            f"No extractable text in {path.name}. This is likely a scanned/"
            "image-only PDF -- OCR is not supported in this build (see README)."
        )
    return text


def _parse_docx(path: Path) -> str:
    try:
        doc = Document(path)
    except Exception as e:
        raise ParseError(f"Could not open DOCX {path.name}: {e}") from e

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Some resumes lay out contact info / skills in tables -- grab those too.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ParseError(f"No extractable text in {path.name}")
    return text


def _parse_txt(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as e:
        raise ParseError(f"Could not read {path.name}: {e}") from e
    if not text:
        raise ParseError(f"{path.name} is empty")
    return text


def find_resume_files(folder: Path) -> list[Path]:
    """Returns supported resume files in a folder, sorted for stable output order."""
    exts = {".pdf", ".docx", ".txt"}
    return sorted(p for p in folder.iterdir() if p.is_file() and p.suffix.lower() in exts)
